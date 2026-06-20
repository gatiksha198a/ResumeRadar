"""
Resume Parser Module
Extracts text and structured information from PDF and DOCX resumes
"""

import io
import re
from typing import Dict, List, Optional, Any
from dataclasses import dataclass, field

import PyPDF2


@dataclass
class ParsedResume:
    """Structured resume data"""
    raw_text: str = ""
    name: str = ""
    email: str = ""
    phone: str = ""
    location: str = ""
    linkedin: str = ""
    github: str = ""
    skills: List[str] = field(default_factory=list)
    education: List[Dict[str, str]] = field(default_factory=list)
    experience: List[Dict[str, Any]] = field(default_factory=list)
    summary: str = ""
    certifications: List[str] = field(default_factory=list)
    languages: List[str] = field(default_factory=list)


class ResumeParser:
    """Parse resumes from PDF and DOCX files"""
    
    # Common skill keywords for extraction
    COMMON_SKILLS = {
        # Programming Languages
        'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'ruby', 'go', 
        'rust', 'scala', 'kotlin', 'swift', 'php', 'perl', 'r', 'matlab',
        # Web Technologies
        'html', 'css', 'react', 'angular', 'vue', 'node.js', 'express', 'django',
        'flask', 'spring', 'asp.net', 'rest api', 'graphql',
        # Data Science & ML
        'machine learning', 'deep learning', 'tensorflow', 'pytorch', 'keras',
        'scikit-learn', 'pandas', 'numpy', 'scipy', 'nlp', 'computer vision',
        'data analysis', 'data visualization', 'tableau', 'power bi',
        # Cloud & DevOps
        'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'jenkins', 'terraform',
        'ansible', 'ci/cd', 'git', 'jira',
        # Databases
        'sql', 'mysql', 'postgresql', 'mongodb', 'redis', 'elasticsearch',
        'oracle', 'sqlite', 'nosql',
        # Soft Skills
        'leadership', 'communication', 'teamwork', 'problem-solving', 
        'project management', 'agile', 'scrum',
        # Other Technical
        'linux', 'unix', 'windows', 'networking', 'security', 'testing',
        'selenium', 'appium', 'junit', 'pytest'
    }
    
    # Education patterns
    DEGREE_PATTERNS = [
        r'\b(Ph\.?D\.?|Doctorate)\b',
        r'\b(M\.?S\.?|M\.?Sc\.?|Master)\b',
        r'\b(B\.?S\.?|B\.?Sc\.?|Bachelor)\b',
        r'\b(M\.?B\.?A\.?)\b',
        r'\b(Associate)\b',
    ]
    
    def __init__(self):
        self.email_pattern = re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b')
        self.phone_pattern = re.compile(r'\b(?:\+?1[-.\s]?)?\(?[0-9]{3}\)?[-.\s]?[0-9]{3}[-.\s]?[0-9]{4}\b')
        self.linkedin_pattern = re.compile(r'linkedin\.com/in/([a-zA-Z0-9-]+)', re.IGNORECASE)
        self.github_pattern = re.compile(r'github\.com/([a-zA-Z0-9-]+)', re.IGNORECASE)
    
    def parse_file(self, file_bytes: bytes, filename: str) -> ParsedResume:
        """Parse resume from file bytes"""
        ext = filename.lower().split('.')[-1]
        
        if ext == 'pdf':
            return self._parse_pdf(file_bytes)
        elif ext == 'docx':
            return self._parse_docx(file_bytes)
        else:
            raise ValueError(f"Unsupported file format: {ext}")
    
    def _parse_pdf(self, file_bytes: bytes) -> ParsedResume:
        """Extract text from PDF"""
        text = ""
        try:
            with io.BytesIO(file_bytes) as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            print(f"Error parsing PDF: {e}")
        
        return self._extract_structured_data(text)
    
    def _parse_docx(self, file_bytes: bytes) -> ParsedResume:
        """Extract text from DOCX"""
        text = ""
        try:
            try:
                from docx import Document
            except ImportError as exc:
                raise RuntimeError(
                    "DOCX parsing is unavailable in this environment because the "
                    "`python-docx` dependency could not load. Please use a PDF file "
                    "or install/allow the required native dependencies."
                ) from exc

            with io.BytesIO(file_bytes) as f:
                doc = Document(f)
                for para in doc.paragraphs:
                    text += para.text + "\n"
                # Also extract tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += cell.text + "\n"
        except Exception as e:
            if isinstance(e, RuntimeError):
                raise
            print(f"Error parsing DOCX: {e}")

        return self._extract_structured_data(text)
    
    def _extract_structured_data(self, text: str) -> ParsedResume:
        """Extract structured information from raw text"""
        resume = ParsedResume(raw_text=text)
        
        # Extract contact info
        resume.email = self._extract_email(text)
        resume.phone = self._extract_phone(text)
        resume.linkedin = self._extract_linkedin(text)
        resume.github = self._extract_github(text)
        
        # Extract name (usually at the top)
        resume.name = self._extract_name(text)
        
        # Extract skills
        resume.skills = self._extract_skills(text)
        
        # Extract education
        resume.education = self._extract_education(text)
        
        # Extract experience
        resume.experience = self._extract_experience(text)
        
        # Extract summary
        resume.summary = self._extract_summary(text)
        
        # Extract certifications
        resume.certifications = self._extract_certifications(text)
        
        return resume
    
    def _extract_email(self, text: str) -> str:
        match = self.email_pattern.search(text)
        return match.group(0) if match else ""
    
    def _extract_phone(self, text: str) -> str:
        match = self.phone_pattern.search(text)
        return match.group(0) if match else ""
    
    def _extract_linkedin(self, text: str) -> str:
        match = self.linkedin_pattern.search(text)
        return match.group(0) if match else ""

    def _extract_github(self, text: str) -> str:
        match = self.github_pattern.search(text)
        return match.group(0) if match else ""
    
    def _extract_name(self, text: str) -> str:
        """Extract name from resume header"""
        lines = text.split('\n')
        for line in lines[:5]:
            line = line.strip()
            # Name is usually short, no special chars, first line
            if 2 <= len(line.split()) <= 4:
                # Check if it looks like a name (no numbers, no email-like chars)
                if not any(c in line for c in ['@', 'http', 'www', '(', ')']):
                    # Check it's not all caps (usually section headers)
                    if line != line.upper():
                        return line
        return ""
    
    def _extract_skills(self, text: str) -> List[str]:
        """Extract skills from text"""
        text_lower = text.lower()
        found_skills = []
        
        for skill in self.COMMON_SKILLS:
            if skill in text_lower:
                # Clean up skill name
                skill_clean = skill.strip()
                if skill_clean not in found_skills:
                    found_skills.append(skill_clean)
        
        return sorted(found_skills)
    
    def _extract_education(self, text: str) -> List[Dict[str, str]]:
        """Extract education details"""
        education = []
        lines = text.split('\n')
        
        # Look for education section
        edu_keywords = ['education', 'academic', 'degree', 'university', 'college']
        
        for i, line in enumerate(lines):
            line_lower = line.lower()
            if any(kw in line_lower for kw in edu_keywords):
                # Extract next 5-10 lines for education details
                for j in range(i+1, min(i+15, len(lines))):
                    edu_line = lines[j].strip()
                    if edu_line:
                        # Try to extract degree and institution
                        degree_match = re.search(r'(Ph\.?D|M\.?S\.?|M\.?Sc|B\.?S\.?|B\.?Sc|M\.?B\.?A|Associate)', edu_line, re.IGNORECASE)
                        year_match = re.search(r'(19|20)\d{2}', edu_line)
                        
                        if degree_match:
                            education.append({
                                'degree': degree_match.group(0),
                                'year': year_match.group(0) if year_match else "",
                                'details': edu_line
                            })
                break
        
        return education
    
    def _extract_experience(self, text: str) -> List[Dict[str, Any]]:
        """Extract work experience"""
        experience = []
        lines = text.split('\n')
        
        # Look for experience section
        exp_keywords = ['experience', 'employment', 'work history', 'professional']
        
        for i, line in enumerate(lines):
            line_lower = line.lower()
            if any(kw in line_lower for kw in exp_keywords):
                # Extract next 20-30 lines for experience details
                for j in range(i+1, min(i+40, len(lines))):
                    exp_line = lines[j].strip()
                    if exp_line and len(exp_line) > 10:
                        # Look for date patterns
                        date_match = re.search(r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s*(19|20)\d{2}', exp_line, re.IGNORECASE)
                        if date_match or exp_line.endswith(':') or exp_line.startswith('-'):
                            experience.append({
                                'text': exp_line,
                                'has_date': bool(date_match)
                            })
                break
        
        return experience[:10]  # Limit to 10 entries
    
    def _extract_summary(self, text: str) -> str:
        """Extract professional summary"""
        lines = text.split('\n')
        
        for i, line in enumerate(lines[:10]):
            line_lower = line.lower()
            if 'summary' in line_lower or 'objective' in line_lower:
                # Get next few lines as summary
                summary_lines = []
                for j in range(i+1, min(i+5, len(lines))):
                    if lines[j].strip():
                        summary_lines.append(lines[j].strip())
                return ' '.join(summary_lines[:3])
        
        # If no summary section, use first paragraph
        return lines[0] if lines else ""
    
    def _extract_certifications(self, text: str) -> List[str]:
        """Extract certifications"""
        certifications = []
        lines = text.split('\n')
        
        cert_keywords = ['certification', 'certificate', 'certified', 'license']
        
        for line in lines:
            line_lower = line.lower()
            if any(kw in line_lower for kw in cert_keywords):
                if line.strip():
                    certifications.append(line.strip())
        
        return certifications[:5]  # Limit to 5
    
    def to_dict(self, resume: ParsedResume) -> Dict[str, Any]:
        """Convert parsed resume to dictionary"""
        return {
            'name': resume.name,
            'email': resume.email,
            'phone': resume.phone,
            'location': resume.location,
            'linkedin': resume.linkedin,
            'github': resume.github,
            'skills': resume.skills,
            'education': resume.education,
            'experience': resume.experience,
            'summary': resume.summary,
            'certifications': resume.certifications,
            'languages': resume.languages,
            'raw_text': resume.raw_text
        }


def parse_resume(file_bytes: bytes, filename: str) -> Dict[str, Any]:
    """Convenience function to parse resume"""
    parser = ResumeParser()
    parsed = parser.parse_file(file_bytes, filename)
    return parser.to_dict(parsed)


if __name__ == "__main__":
    # Test with sample
    print("Resume Parser Module Loaded")
    print(f"Common skills tracked: {len(ResumeParser.COMMON_SKILLS)}")
